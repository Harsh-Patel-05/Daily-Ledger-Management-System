"""
Generate Daily Ledger Management System — Project Documentation (.docx)
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "docs" / "screenshots"
OUT = ROOT / "docs" / "Daily_Ledger_Management_System_Documentation.docx"

# Prefer fresh captures if present, else fall back to existing assets
SHOT_MAP = {
    "login": [
        "01_login.png",
        "image-5b1998f3-52b8-4ec0-9982-ab7bd072281d.png",
    ],
    "signup": [
        "02_signup.png",
        "image-fec670e7-0809-4c13-9bfe-9936c8ca150f.png",
    ],
    "dashboard": [
        "03_dashboard.png",
        "image-ee376d3f-847b-4fee-9e89-085ef3942924.png",
    ],
    "tour": [
        "04_tour.png",
        "image-ee376d3f-847b-4fee-9e89-085ef3942924.png",
    ],
    "customers": ["05_customers.png"],
    "transactions": [
        "06_transactions.png",
        "image-a861bf7b-2fee-49b4-8a6c-323c0d6b356a.png",
    ],
    "invoices": [
        "07_invoices.png",
        "image-83aeade1-ae07-486d-8462-1f835ea23be3.png",
    ],
    "create_invoice": [
        "08_create_invoice.png",
        "image-5308a484-a324-46b1-8c5e-425da339d551.png",
    ],
    "ledger": ["09_ledger.png"],
    "reports": ["10_reports.png"],
    "analytics": ["11_analytics.png"],
    "notifications": [
        "12_notifications.png",
        "image-96d9bd23-dedd-4746-84bf-530e2b6a7b2d.png",
    ],
    "profile": [
        "13_profile.png",
        "image-52dfbee1-ffb7-4a9c-b147-1e0d49864f71.png",
    ],
    "settings": [
        "14_settings.png",
        "image-cb4b0188-eda5-49e1-b2fc-e7ca78c86877.png",
    ],
    "alerts": [
        "15_alert_settings.png",
        "image-12341bfb-7015-47c9-a11e-a73ce94cf3bf.png",
    ],
    "reminder": [
        "16_reminder.png",
        "image-3fae3739-22c2-47ab-a5c2-f05b7e1320b9.png",
    ],
}


def find_shot(key: str):
    for name in SHOT_MAP.get(key, []):
        p = SHOTS / name
        if p.exists():
            return p
    return None


def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return h


def add_para(doc, text, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)


def add_screenshot(doc, key, caption):
    path = find_shot(key)
    if not path:
        add_para(doc, f"[Screenshot: {caption} — capture pending]", size=10, bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Keep images readable in Word
    run.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure: {caption}")
    set_run_font(r, size=9, bold=True, color=(0x47, 0x55, 0x69))
    cap.paragraph_format.space_after = Pt(14)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            for p in cells[c_i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def build():
    SHOTS.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ========== COVER ==========
    for _ in range(2):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Daily Ledger Management System")
    set_run_font(r, size=28, bold=True, color=(0x1E, 0x3A, 0x8A))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Complete Project Documentation")
    set_run_font(r, size=16, bold=True, color=(0x33, 0x41, 0x55))

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("Digital Roj Mel for Indian Shopkeepers")
    set_run_font(r, size=12, color=(0x64, 0x74, 0x8B))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Version 1.0  ·  Document Date: {date.today().strftime('%d %B %Y')}\n"
        "Frontend: React 19 + Vite  ·  Backend: Django REST + JWT"
    )
    set_run_font(r, size=10, color=(0x64, 0x74, 0x8B))

    doc.add_page_break()

    # ========== TOC-like overview ==========
    add_heading_styled(doc, "1. Table of Contents", 1)
    add_bullets(doc, [
        "1. Introduction & Purpose",
        "2. Technology Stack",
        "3. System Architecture",
        "4. Getting Started (Setup & Login)",
        "5. Module-wise Features (with Screenshots)",
        "6. Quick Actions, Modals & Website Tour",
        "7. Theme Customization",
        "8. Backend API Overview",
        "9. Security & Authentication",
        "10. Demo Credentials & Keyboard Shortcuts",
        "11. Project Structure",
        "12. Conclusion",
    ])

    # ========== 1 ==========
    add_heading_styled(doc, "2. Introduction & Purpose", 1)
    add_para(
        doc,
        "Daily Ledger Management System is a full-stack web application designed as a digital "
        "Roj Mel (daily shop ledger) for Indian shopkeepers and traders. It helps shops track "
        "customer credit (udhaar), receive collections, generate GST-style invoices, manage "
        "expenses, send payment reminders, and close the day with a clear cash summary.",
    )
    add_para(doc, "Core business goals:", bold=True)
    add_bullets(doc, [
        "Replace paper-based daily ledgers with a reliable digital system",
        "Track customer balances, credit limits, and overdue accounts",
        "Record credit sales and payments (including partial payments)",
        "Create, share, and collect against tax invoices",
        "Provide dashboards, reports, and analytics for business decisions",
        "Support multi-channel owner alerts (in-app, email, SMS) and customer reminders",
    ])

    # ========== 2 ==========
    add_heading_styled(doc, "3. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technologies"],
        [
            ["Frontend", "React 19, Vite 8, React Router 7, Tailwind CSS v4, Framer Motion, Recharts, Context API"],
            ["Backend", "Django 6, Django REST Framework, SimpleJWT, django-filter, CORS, drf-spectacular"],
            ["Database", "SQLite (development) — ready for PostgreSQL/MySQL in production"],
            ["Auth", "JWT access + refresh tokens"],
            ["Extras", "Tesseract.js OCR for invoice upload, html2canvas + jsPDF for invoice PDF"],
        ],
    )

    # ========== 3 ==========
    add_heading_styled(doc, "4. System Architecture", 1)
    add_para(
        doc,
        "The React SPA talks to the Django REST API under /api. In development, Vite proxies "
        "/api and /media to http://127.0.0.1:8000. After login, JWT tokens are stored in the "
        "browser and attached to authenticated requests. App state (customers, transactions, "
        "invoices, settings) is loaded from the API into React Context.",
    )
    add_bullets(doc, [
        "Frontend (port 5173) → Vite proxy → Backend API (port 8000)",
        "Auth endpoints under /api/auth/",
        "Business modules: customers, transactions, invoices, notifications, core (dashboard/ledger/reports/analytics)",
        "Media files (logos) served from Django /media/",
    ])

    # ========== 4 ==========
    add_heading_styled(doc, "5. Getting Started", 1)
    add_heading_styled(doc, "5.1 Backend setup", 2)
    add_para(doc, "From the backend folder:")
    add_bullets(doc, [
        "python -m venv venv && activate the virtual environment",
        "pip install -r requirements.txt",
        "copy .env.example to .env and configure secrets if needed",
        "python manage.py migrate",
        "python manage.py seed_demo  (optional demo data)",
        "python manage.py runserver 8000",
    ])
    add_heading_styled(doc, "5.2 Frontend setup", 2)
    add_bullets(doc, [
        "npm install",
        "npm run dev",
        "Open http://localhost:5173",
    ])
    add_heading_styled(doc, "5.3 Login & Sign up", 2)
    add_para(
        doc,
        "The auth screens use a split layout: branding on the left and form on the right. "
        "Users can register a new shop account or sign in with existing credentials.",
    )
    add_screenshot(doc, "login", "Login page — Welcome back / Sign In")
    add_screenshot(doc, "signup", "Sign up page — Create account")

    # ========== 5 Modules ==========
    add_heading_styled(doc, "6. Module-wise Features (with Screenshots)", 1)

    add_heading_styled(doc, "6.1 Dashboard", 2)
    add_para(
        doc,
        "The dashboard is the home screen after login. It shows live KPIs (today’s sales, "
        "collection, pending amount, invoice due, customers, transactions, unpaid invoices, "
        "overdue customers), monthly collection charts, credit vs paid breakdown, unpaid "
        "invoices, activity log, and alerts. Quick buttons open Collect Payment and Due lists.",
    )
    add_screenshot(doc, "dashboard", "Dashboard — KPIs, charts, and live overview")

    add_heading_styled(doc, "6.2 Customers", 2)
    add_para(
        doc,
        "Manage customer master data: name, business, mobile, GST, address, credit limit, "
        "and status (active / overdue). Customer details show utilization, ledger shortcuts, "
        "and record-payment actions.",
    )
    add_screenshot(doc, "customers", "Customers module")

    add_heading_styled(doc, "6.3 Transactions", 2)
    add_para(
        doc,
        "Daily ledger entries support Credit (Maal Diya), Payment Received, Return, Discount, "
        "and Expense. Forms capture date, customer, item description, qty/rate/amount, and "
        "payment method. Credit-limit warnings appear when a sale would exceed the limit.",
    )
    add_screenshot(doc, "transactions", "New Transaction form")

    add_heading_styled(doc, "6.4 Invoices", 2)
    add_para(
        doc,
        "Create GST-style invoices with line items, discount, tax rate, notes, and terms. "
        "Features include multiple formats, PDF download/print, share, duplicate, mark paid "
        "(syncs ledger payment), record payment against an invoice, and OCR upload.",
    )
    add_screenshot(doc, "invoices", "Invoices list")
    add_screenshot(doc, "create_invoice", "Create Invoice screen with line items and totals")

    add_heading_styled(doc, "6.5 Ledger", 2)
    add_para(
        doc,
        "Per-customer running balance with timeline/table views. Useful before collecting "
        "dues or reviewing a customer’s full khata.",
    )
    add_screenshot(doc, "ledger", "Customer Ledger")

    add_heading_styled(doc, "6.6 Reports & Analytics", 2)
    add_para(
        doc,
        "Reports summarize credit, collection, expenses, and debtors for a period. "
        "Analytics shows multi-month trends using charts.",
    )
    add_screenshot(doc, "reports", "Reports")
    add_screenshot(doc, "analytics", "Analytics")

    add_heading_styled(doc, "6.7 Notifications", 2)
    add_para(
        doc,
        "In-app alerts for payment due, overdue, daily summary, and invoice reminders. "
        "Users can refresh, filter by type, mark as read, and open customer reminders.",
    )
    add_screenshot(doc, "notifications", "Notifications list")
    add_screenshot(doc, "reminder", "Send Payment Reminder modal")

    add_heading_styled(doc, "6.8 Profile", 2)
    add_para(
        doc,
        "Shop owner profile: personal/business details, GST, invoice prefix, address, "
        "and shop logo used on invoices.",
    )
    add_screenshot(doc, "profile", "Profile — business details and shop logo")

    add_heading_styled(doc, "6.9 Settings", 2)
    add_para(
        doc,
        "Business information, bank & UPI details (printed on invoices), preferences "
        "(currency, language, fiscal year), appearance (theme mode + color theme), "
        "owner alert toggles, backup/export, and website tour restart.",
    )
    add_screenshot(doc, "settings", "Settings — business and banking configuration")
    add_screenshot(doc, "alerts", "Owner alert / notification preference toggles")

    # ========== 6 Tour & Quick actions ==========
    add_heading_styled(doc, "7. Quick Actions, Modals & Website Tour", 1)
    add_para(doc, "Navbar quick actions (desktop) and mobile + menu include:", bold=True)
    add_bullets(doc, [
        "Record Payment — collect from customers with 25% / 50% / Full chips",
        "Collections Due — overdue and pending customers + invoice due list",
        "Day Closing (Roj Mel) — today’s credit, collection, expense, method-wise summary (saved to server)",
        "Quick Expense — shop expenses by category",
        "Quick Customer / Transaction / Invoice",
        "Command Palette (Ctrl+K) for search and actions",
    ])
    add_para(
        doc,
        "Website Tour: first login auto-starts an 11-step guided tour. Users can replay "
        "it from the help (?) icon, profile menu, Settings → Help & Tour, or Ctrl+K → "
        "Start Website Tour. Content is in clear English with polished welcome/finish cards.",
    )
    add_screenshot(doc, "tour", "Website tour overlay on the dashboard")

    # ========== 7 Theme ==========
    add_heading_styled(doc, "8. Theme Customization", 1)
    add_para(
        doc,
        "Appearance settings support Light, Dark, and System modes, plus full color themes "
        "(Blue, Emerald, Teal, Indigo, Violet, Rose, Orange, Amber, Cyan, Slate). Selecting "
        "a color retints primary actions, backgrounds, surfaces, borders, and slate tones "
        "across the app. Preferences persist in localStorage and BusinessSettings (theme + accentColor).",
    )

    # ========== 8 API ==========
    add_heading_styled(doc, "9. Backend API Overview", 1)
    add_table(
        doc,
        ["Area", "Base path", "Key operations"],
        [
            ["Auth", "/api/auth/", "login, register, refresh, profile, settings, OTP reset"],
            ["Customers", "/api/customers/", "CRUD, recalculate balance"],
            ["Transactions", "/api/transactions/", "CRUD, record-payment, summary, day-close"],
            ["Invoices", "/api/invoices/", "CRUD, next-number, duplicate, mark-paid"],
            ["Notifications", "/api/notifications/", "list, sync, reminders, activity, alerts"],
            ["Core", "/api/", "dashboard, ledger, reports, analytics, health"],
        ],
    )
    add_para(doc, "Interactive Swagger docs: http://127.0.0.1:8000/api/docs/")

    # ========== 9 Security ==========
    add_heading_styled(doc, "10. Security & Authentication", 1)
    add_bullets(doc, [
        "JWT Bearer authentication on protected endpoints",
        "Refresh token rotation support via /api/auth/refresh/",
        "Password reset via email OTP flow",
        "CORS configured for local development and ngrok hosts",
        "Owner-only data scoping (each user sees only their shop records)",
    ])

    # ========== 10 Demo ==========
    add_heading_styled(doc, "11. Demo Credentials & Shortcuts", 1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Demo Email", "mukesh@ganeshtraders.com"],
            ["Demo Password", "password123"],
            ["Frontend", "http://localhost:5173"],
            ["Backend API", "http://127.0.0.1:8000/api/"],
            ["API Docs", "http://127.0.0.1:8000/api/docs/"],
            ["Command Palette", "Ctrl+K / ⌘K"],
            ["Website Tour", "Navbar ? icon or Settings → Start Tour"],
        ],
    )

    # ========== 11 Structure ==========
    add_heading_styled(doc, "12. Project Structure", 1)
    add_para(doc, "Frontend (src/):", bold=True)
    add_bullets(doc, [
        "pages/ — Dashboard, Customers, Transactions, Invoices, Ledger, Reports, Analytics, Notifications, Profile, Settings, Auth",
        "components/ — UI kit, layout (Navbar/Sidebar), modals, payments, tour, invoice templates",
        "context/ — Auth, App, Theme, Toast, Modal, Tour",
        "api/ — REST client modules for all backend domains",
        "data/ — defaults, theme presets, tour steps, invoice formats",
    ])
    add_para(doc, "Backend (backend/):", bold=True)
    add_bullets(doc, [
        "accounts/ — users, profile, business settings, OTP",
        "customers/, transactions/, invoices/, notifications/, core/",
        "config/ — Django settings, URLs, JWT, CORS",
    ])

    # ========== 12 Conclusion ==========
    add_heading_styled(doc, "13. Conclusion", 1)
    add_para(
        doc,
        "Daily Ledger Management System delivers an end-to-end digital ledger for shops: "
        "from login and onboarding tour, through daily credit/collection workflows, invoices, "
        "reminders, day closing, theming, and reporting — backed by a proper Django REST API "
        "and a modern React frontend. This document, together with the embedded screenshots, "
        "serves as a complete reference for users, testers, and developers.",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("— End of Document —")
    set_run_font(r, size=10, color=(0x64, 0x74, 0x8B))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")
    return OUT


if __name__ == "__main__":
    build()
